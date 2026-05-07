"""Convert Markdown report to Word document."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re, os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, align=None, size=12, italic=False):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_table_from_data(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    doc.add_paragraph()

# ---- COVER PAGE ----
doc.add_paragraph()
add_para("Kolhapur Institute of Technology's", bold=True, align='center', size=16)
add_para("College of Engineering (Empowered Autonomous)", bold=True, align='center', size=14)
add_para("Kolhapur", bold=True, align='center', size=14)
doc.add_paragraph()
add_para('"ThinkFlow: AI-Powered Real-Time Factory Safety Monitoring System"', bold=True, align='center', size=15)
doc.add_paragraph()
add_para("Submitted in partial fulfillment of the requirements of the degree of", align='center', size=12)
doc.add_paragraph()
add_para("Mini Project-IV", bold=True, align='center', size=14)
add_para("UCBIL0671", bold=True, align='center', size=13)
doc.add_paragraph()
add_para("By", bold=True, align='center', size=12)
add_table_from_data(["Roll No.", "Name", "PRN"], [["___", "_______________", "_______________"], ["___", "_______________", "_______________"], ["___", "_______________", "_______________"]])
add_para("Under the Supervision of", bold=True, align='center', size=12)
add_para("Mrs. Yashaswini A. Kadiyal", bold=True, align='center', size=13)
add_para("Asst. Professor", align='center', size=12)
add_para("Department of CSBS", align='center', size=12)
doc.add_paragraph()
add_para("Department of Computer Science and Business Systems", bold=True, align='center', size=12)
add_para("KITCOE, Kolhapur", bold=True, align='center', size=12)
add_para("May, 2026", bold=True, align='center', size=12)
doc.add_page_break()

# ---- CERTIFICATE ----
add_heading_styled("CERTIFICATE", level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('This is to certify that the following students of Second Year Computer Science and Business Systems has submitted Mini Project-IV report entitled "ThinkFlow: AI-Powered Real-Time Factory Safety Monitoring System" in partial fulfillment for the course code "UCBIL0671", in the academic year 2025-26. It has been found to be satisfactory and hereby approved for submission.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
add_table_from_data(["Sr. No.", "Name", "PRN"], [["1", "_______________", "_______________"], ["2", "_______________", "_______________"], ["3", "_______________", "_______________"]])
add_para("Guide Name: Mrs. Yashaswini Kadiyal", bold=True, size=12)
add_para("Asst. Prof., CSBS", size=12)
doc.add_page_break()

# ---- ACKNOWLEDGEMENT ----
add_heading_styled("ACKNOWLEDGEMENT", level=1)
ack = """We would like to express our sincere gratitude to our project guide Mrs. Yashaswini A. Kadiyal, Assistant Professor, Department of Computer Science and Business Systems, KITCOE Kolhapur, for her invaluable guidance, constant encouragement, and support throughout the development of this project.

We are also grateful to the Head of Department of CSBS for providing us the opportunity and resources to work on this project. We thank the entire faculty of the CSBS department for their knowledge and motivation.

We extend our thanks to KIT's College of Engineering, Kolhapur for providing the necessary infrastructure, lab facilities, and computing resources.

Finally, we thank our family and friends for their continuous encouragement during this academic endeavor."""
for para in ack.split('\n\n'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(para.strip())
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
doc.add_page_break()

# ---- ABSTRACT ----
add_heading_styled("ABSTRACT", level=1)
abstract = """Industrial safety remains a critical concern in manufacturing environments worldwide. Traditional safety monitoring approaches rely heavily on human supervision, which is prone to fatigue, inattention, and delayed response times. This project presents ThinkFlow, an AI-powered real-time factory safety monitoring system that leverages computer vision and deep learning to automatically detect human presence in designated danger zones and trigger instant alerts.

The system is built using a microservices architecture comprising three independent services: (1) a Python AI Service utilizing YOLOv8 for real-time human detection and ByteTrack for persistent object tracking, (2) a Java Spring Boot Backend serving as the API gateway with JWT authentication, WebSocket communication, danger zone logic, and SQLite database management, and (3) a React Frontend providing a modern, dark-themed dashboard for real-time monitoring, camera management, zone configuration, and event history.

Key features include multi-camera support, polygon-based danger zone definition with ray-casting breach detection, a 30-second alert cooldown system, STOMP WebSocket for instant alert delivery, and audible sound notifications using the Web Audio API. The system operates fully offline on a local network with no cloud dependency, making it suitable for secure industrial environments.

Testing demonstrated that the system successfully detects human presence with high accuracy (~85-95% confidence), maintains persistent tracking IDs across frames, and delivers zone breach alerts within 200ms of detection. The modular architecture enables easy extension for future capabilities such as PPE detection, sensor integration, and multi-factory deployment."""
for para in abstract.split('\n\n'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(para.strip())
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
add_para("Keywords: Computer Vision, YOLOv8, Object Detection, Factory Safety, Real-Time Monitoring, Microservices, WebSocket, Spring Boot, React, Deep Learning", bold=True, italic=True, size=11)
doc.add_page_break()

# ---- Now read the full markdown and parse remaining chapters ----
md_path = r"c:\Users\rushi\OneDrive\Desktop\factory_safety_system\ThinkFlow_Mini_Project_Report_FULL.md"
with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

# Find chapters starting from CHAPTER 1
chapters = re.split(r'\n# (CHAPTER \d+)', content)

for i in range(1, len(chapters), 2):
    chapter_title = chapters[i]
    chapter_body = chapters[i + 1] if i + 1 < len(chapters) else ""
    
    # Get full title from first line
    lines = chapter_body.strip().split('\n')
    full_title = chapter_title
    if lines and lines[0].startswith(':'):
        full_title = chapter_title + lines[0]
        lines = lines[1:]
    
    # Clean title
    full_title = full_title.replace(':', '').strip()
    add_heading_styled(full_title, level=1)
    
    body = '\n'.join(lines)
    
    # Process line by line
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Skip markdown artifacts like --- or ```
        if line.startswith('---'):
            doc.add_page_break()
            continue
        if line.startswith('```'):
            continue
            
        # Sub-headings
        if line.startswith('## '):
            text = line.replace('## ', '').replace('**', '')
            add_heading_styled(text, level=2)
        elif line.startswith('### '):
            text = line.replace('### ', '').replace('**', '')
            add_heading_styled(text, level=3)
        # Table rows
        elif line.startswith('|') and not line.startswith('|--'):
            # Skip tables for now, handled separately
            pass
        # Numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            p = doc.add_paragraph(text, style='List Number')
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(12)
        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])
            p = doc.add_paragraph(text, style='List Bullet')
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(12)
        # Fig references
        elif line.startswith('Fig ') or line.startswith('*[INSERT'):
            add_para(line.replace('*', ''), italic=True, align='center', size=11)
        # Regular text
        elif not line.startswith('┌') and not line.startswith('│') and not line.startswith('└') and not line.startswith('▼') and not line.startswith('▲') and not line.startswith('+'):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            text = text.replace('`', '')
            if text.strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                r = p.add_run(text)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)

# ---- REFERENCES (manually add for proper formatting) ----
doc.add_page_break()
add_heading_styled("REFERENCES", level=1)

refs = [
    '[1] International Labour Organization (ILO), "Safety and Health at Work," ILO Global Estimates, 2023.',
    '[2] National Crime Records Bureau (NCRB), "Accidental Deaths and Suicides in India," Ministry of Home Affairs, Government of India, 2023.',
    '[3] D. Keval and M. A. Sasse, "Not the Usual Suspects: A Study of Factors Reducing the Effectiveness of CCTV," Security Journal, vol. 23, no. 2, pp. 134-154, 2010.',
    '[4] J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, "You Only Look Once: Unified, Real-Time Object Detection," IEEE CVPR, pp. 779-788, 2016.',
    '[5] J. Redmon and A. Farhadi, "YOLO9000: Better, Faster, Stronger," IEEE CVPR, pp. 6517-6525, 2017.',
    '[6] A. Bochkovskiy, C.-Y. Wang, and H.-Y. M. Liao, "YOLOv4: Optimal Speed and Accuracy of Object Detection," arXiv:2004.10934, 2020.',
    '[7] G. Jocher, A. Chaurasia, and J. Qiu, "Ultralytics YOLOv8," Ultralytics, 2023.',
    '[8] M. Zhang et al., "Deep Learning-Based Worker Detection for Construction Site Safety Monitoring," Automation in Construction, vol. 128, 2021.',
    '[9] N. D. Nath, A. H. Behzadan, and S. G. Paal, "Deep Learning for Site Safety: Real-Time Detection of PPE," Automation in Construction, vol. 112, 2020.',
    '[10] W. Fang et al., "Falls from Heights: A Computer Vision-Based Approach for Safety Harness Detection," Automation in Construction, vol. 91, pp. 53-61, 2018.',
    '[11] Y. Zhang et al., "ByteTrack: Multi-Object Tracking by Associating Every Detection Box," ECCV, pp. 1-21, 2022.',
    '[12] N. Wojke, A. Bewley, and D. Paulus, "Simple Online and Realtime Tracking with a Deep Association Metric," IEEE ICIP, pp. 3645-3649, 2017.',
    '[13] N. Dragoni et al., "Microservices: Yesterday, Today, and Tomorrow," Present and Ulterior Software Engineering, Springer, pp. 195-216, 2017.',
    '[14] S. Newman, "Building Microservices: Designing Fine-Grained Systems," O\'Reilly Media, 2nd Ed., 2021.',
    '[15] I. Fette and A. Melnikov, "The WebSocket Protocol," RFC 6455, IETF, 2011.',
    '[16] STOMP Protocol Specification, "Simple Text Oriented Messaging Protocol," Version 1.2, 2012.',
    '[17] Spring Framework, "Spring Boot Reference Documentation," Pivotal Software, Version 3.2.2, 2024.',
    '[18] Meta Platforms, "React Documentation," 2024.',
    '[19] D. P. Kingma and J. Ba, "Adam: A Method for Stochastic Optimization," ICLR, 2015.',
    '[20] T. Lin et al., "Microsoft COCO: Common Objects in Context," ECCV, pp. 740-755, 2014.',
    '[21] OpenCV Library, "Open Source Computer Vision Library," 2024.',
    '[22] S. Tiangolo, "FastAPI: Modern, Fast Web Framework for Building APIs," 2024.',
    '[23] Auth0, "JSON Web Tokens (JWT) Introduction," 2024.',
    '[24] SQLite Consortium, "SQLite: A Self-Contained, Serverless SQL Database Engine," 2024.',
    '[25] E. You, "Vite: Next Generation Frontend Tooling," 2024.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

# Save
out = r"c:\Users\rushi\OneDrive\Desktop\factory_safety_system\ThinkFlow_Mini_Project_Report.docx"
doc.save(out)
print(f"Word document saved: {out}")
